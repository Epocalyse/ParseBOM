import docx
import pandas as pd
import re
import os
from ExtractBOM import ExtractBOM


class Convert:
    @staticmethod
    def bomDocToCSV(doc, loc):
        # TODO:Change docs to docx
        # doc = docx.Document('BOMWordCopies/102024GMP.docx')
        docu = docx.Document(loc + '/' + doc)
        table = docu.tables[0]

        # Data will be a list of rows represented as tuples containing each row's data.
        data = []

        for i, row in enumerate(table.rows):
            text = (cell.text for cell in row.cells)

            # Establish the mapping based on the first row
            # headers; these will become the keys of our dictionary
            if i == 0:
                continue

            # Construct a tuple for row
            row_data = tuple(text)
            data.append(row_data)

        df = pd.DataFrame(data)

        # Removing first two rows
        df = df.iloc[2:]
        # Appending header names
        df.columns = ['MaterialType', 'CultivationName', 'SupplierCatalogueNumber', 'ERPnumber', 'ERPInventoryControl',
                      'LIMSSpec', 'BOM']

        extraction = ExtractBOM()
        df.to_csv('BOMNewCopies/' + extraction.renameBOM(doc) + "BOM.csv", mode='w', header=True)

    @staticmethod
    def sopDocToCSV(doc, loc):
        cells = ['UK SKU', 'SG SKU', 'SKU', 'Supplier']
        equipment_data = []
        bom_data = []
        record_equipment = False
        record_material = False
        record_reagent = False

        for para in doc.paragraphs:
            if para.text == 'PROCEDURE':
                break
            elif para.text == 'Critical reagents' or record_reagent:
                record_reagent = True

                text = re.split(', |\(', para.text)
                clean_text = [text[0]]
                clean_text.extend(word for word in text if any(map(word.__contains__, cells)))

                clean_entry = [text[0]]
                for tex in text:
                    for cell in cells:
                        if any(map(tex.__contains__, cells)):
                            if tex.__contains__("UK SKU"):
                                clean_entry.append(re.split('UK SKU |\)', tex)[1])
                            if tex.__contains__("SG SKU"):
                                # print(re.split('SG SKU:|\s', text)[1])
                                clean_entry.append(re.split('SG SKU:|\s', tex)[1])
                            if tex.__contains__("SKU"):
                                pass
                                # cleanEntry.extend(text)
                            if tex.__contains__("Supplier"):
                                pass
                                # cleanEntry.extend(text)
                        else:
                            # Currently having issues
                            clean_entry.append("")
                # print("AAAAAAAAAAAAAAAAAAA_________________")
                # print(cleanEntry)
                # print(cleanText)

                row_data = tuple(clean_text)
                bom_data.append(row_data)
            elif para.text == 'Equipment':
                record_equipment = True
            elif record_equipment:
                equipment_data.append((para.text, 'Equipment'))
            elif para.text == 'Materials' or record_material:
                record_material = True
                equipment_data.append((para.text, 'Material'))

        ef = pd.DataFrame(equipment_data)
        df = pd.DataFrame(bom_data)

        # df = df.iloc[1:]
        # df.columns = ['MaterialType', 'UKSKU', 'SGSKU']

        ef.to_csv('BOMNewCopies/BPBT3107_Equipment.csv', mode='w', header=False)
        df.to_csv('BOMNewCopies/BPBT3107_BOM.csv', mode='w', header=False)


# conv = Convert()
# conv.bomDocToCSV('102024GMP.docx', 'BOMWordCopies')
